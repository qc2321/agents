"""
Resume RAG (Retrieval-Augmented Generation) System
Processes Word documents and creates a searchable knowledge base using Chroma DB
"""

import os
import chromadb
from chromadb.config import Settings
from docx import Document
from openai import OpenAI
from typing import List, Dict, Any
import json
import uuid
import re


class ResumeRAG:
    def __init__(self, db_path: str = "./chroma_db", collection_name: str = "resume_knowledge"):
        """
        Initialize the Resume RAG system
        
        Args:
            db_path: Path to store the Chroma database
            collection_name: Name of the collection to store resume data
        """
        self.openai = OpenAI()
        self.db_path = db_path
        self.collection_name = collection_name
        
        # Initialize Chroma client
        self.client = chromadb.PersistentClient(path=db_path)
        
        # Get or create collection
        try:
            self.collection = self.client.get_collection(name=collection_name)
            print(f"Loaded existing collection: {collection_name}")
        except:
            self.collection = self.client.create_collection(
                name=collection_name,
                metadata={"description": "Resume and career information knowledge base"}
            )
            print(f"Created new collection: {collection_name}")
    
    def extract_text_from_docx(self, docx_path: str) -> Dict[str, Any]:
        """
        Extract text content from a Word document
        
        Args:
            docx_path: Path to the .docx file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            doc = Document(docx_path)
            
            # Extract all paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # Only add non-empty paragraphs
                    paragraphs.append(text)
            
            # Extract tables
            tables_content = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))
                if table_text:
                    tables_content.append("\n".join(table_text))
            
            return {
                "paragraphs": paragraphs,
                "tables": tables_content,
                "full_text": "\n".join(paragraphs + tables_content),
                "filename": os.path.basename(docx_path)
            }
            
        except Exception as e:
            print(f"Error extracting text from {docx_path}: {e}")
            return {"paragraphs": [], "tables": [], "full_text": "", "filename": ""}
    
    def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 100) -> List[str]:
        """
        Split text into overlapping chunks for better retrieval
        
        Args:
            text: Text to chunk
            chunk_size: Maximum characters per chunk
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence or paragraph boundaries
            if end < len(text):
                # Look for sentence endings
                last_period = text.rfind('.', start, end)
                last_newline = text.rfind('\n', start, end)
                
                if last_period > start + chunk_size // 2:
                    end = last_period + 1
                elif last_newline > start + chunk_size // 2:
                    end = last_newline + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
            if start >= len(text):
                break
        
        return chunks
    
    def get_embedding(self, text: str) -> List[float]:
        """
        Get OpenAI embedding for text
        
        Args:
            text: Text to embed
            
        Returns:
            Embedding vector
        """
        try:
            response = self.openai.embeddings.create(
                model="text-embedding-3-small",
                input=text
            )
            return response.data[0].embedding
        except Exception as e:
            print(f"Error getting embedding: {e}")
            return []
    
    def add_resume(self, docx_path: str) -> bool:
        """
        Process and add a resume to the knowledge base
        
        Args:
            docx_path: Path to the resume Word document
            
        Returns:
            True if successful, False otherwise
        """
        print(f"Processing resume: {docx_path}")
        
        # Extract text from document
        doc_data = self.extract_text_from_docx(docx_path)
        
        if not doc_data["full_text"]:
            print("No text extracted from document")
            return False
        
        # Create chunks from the full text
        chunks = self.chunk_text(doc_data["full_text"])
        print(f"Created {len(chunks)} chunks from resume")
        
        # Process each chunk
        documents = []
        metadatas = []
        ids = []
        embeddings = []
        
        for i, chunk in enumerate(chunks):
            # Generate embedding
            embedding = self.get_embedding(chunk)
            if not embedding:
                continue
            
            # Create unique ID
            chunk_id = f"{doc_data['filename']}_chunk_{i}"
            
            # Prepare metadata
            metadata = {
                "source": doc_data["filename"],
                "chunk_index": i,
                "chunk_type": "resume_content",
                "document_type": "resume"
            }
            
            documents.append(chunk)
            metadatas.append(metadata)
            ids.append(chunk_id)
            embeddings.append(embedding)
        
        # Add to collection
        try:
            self.collection.add(
                documents=documents,
                metadatas=metadatas,
                ids=ids,
                embeddings=embeddings
            )
            print(f"Successfully added {len(documents)} chunks to knowledge base")
            return True
        except Exception as e:
            print(f"Error adding to collection: {e}")
            return False
    
    def search(self, query: str, n_results: int = 5) -> List[Dict[str, Any]]:
        """
        Search the resume knowledge base
        
        Args:
            query: Search query
            n_results: Number of results to return
            
        Returns:
            List of search results with content and metadata
        """
        try:
            # Get embedding for query
            query_embedding = self.get_embedding(query)
            if not query_embedding:
                return []
            
            # Search the collection
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=n_results,
                include=['documents', 'metadatas', 'distances']
            )
            
            # Format results
            formatted_results = []
            if results['documents'] and results['documents'][0]:
                for i in range(len(results['documents'][0])):
                    formatted_results.append({
                        'content': results['documents'][0][i],
                        'metadata': results['metadatas'][0][i],
                        'distance': results['distances'][0][i]
                    })
            
            return formatted_results
            
        except Exception as e:
            print(f"Error searching: {e}")
            return []
    
    def get_relevant_context(self, query: str, max_context_length: int = 2000) -> str:
        """
        Get relevant context for a query, formatted for use in prompts
        
        Args:
            query: The user's question/query
            max_context_length: Maximum length of context to return
            
        Returns:
            Formatted context string
        """
        results = self.search(query, n_results=5)
        
        if not results:
            return ""
        
        context_parts = []
        current_length = 0
        
        for result in results:
            content = result['content']
            if current_length + len(content) <= max_context_length:
                context_parts.append(content)
                current_length += len(content)
            else:
                # Add partial content if there's space
                remaining_space = max_context_length - current_length
                if remaining_space > 100:  # Only add if meaningful space left
                    context_parts.append(content[:remaining_space] + "...")
                break
        
        if context_parts:
            return "\n\n---\n\n".join(context_parts)
        return ""
    
    def list_collections(self) -> List[str]:
        """List all collections in the database"""
        return [collection.name for collection in self.client.list_collections()]
    
    def get_collection_info(self) -> Dict[str, Any]:
        """Get information about the current collection"""
        try:
            count = self.collection.count()
            return {
                "name": self.collection_name,
                "count": count,
                "path": self.db_path
            }
        except Exception as e:
            return {"error": str(e)}
    
    def delete_collection(self):
        """Delete the current collection (use with caution!)"""
        try:
            self.client.delete_collection(name=self.collection_name)
            print(f"Deleted collection: {self.collection_name}")
        except Exception as e:
            print(f"Error deleting collection: {e}")


# Example usage and testing functions
def test_resume_rag():
    """Test the ResumeRAG system"""
    rag = ResumeRAG()
    
    print("Collection info:", rag.get_collection_info())
    
    # Test search
    test_queries = [
        "What is my work experience?",
        "What are my technical skills?",
        "Tell me about my education",
        "What projects have I worked on?"
    ]
    
    for query in test_queries:
        print(f"\nQuery: {query}")
        results = rag.search(query, n_results=3)
        for i, result in enumerate(results):
            print(f"  Result {i+1}: {result['content'][:100]}...")
            print(f"  Distance: {result['distance']:.3f}")


if __name__ == "__main__":
    # Example usage
    rag = ResumeRAG()
    
    # Add a resume (you'll need to provide the path to your Word document)
    # rag.add_resume("path/to/your/resume.docx")
    
    # Test the system
    test_resume_rag()
